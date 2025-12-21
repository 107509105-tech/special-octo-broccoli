# translate_deepl_bilingual.py
# 使用 deep_translator 進行雙語處理（保留中文，下方添加英文）

import os
os.system('cls' if os.name == 'nt' else 'clear')

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import time
from pathlib import Path
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
from deep_translator import GoogleTranslator

#! 固定對照表讀取 get_fixed_or_translator
#! 新增英文行 add_english_below
#! 調整字體大小 set_paragraph_font_size

# ==================== 設定區 ====================
# Deep Translator 設定
translator = GoogleTranslator(source='zh-TW', target='en')

# 固定翻譯對照表
FIXED_MAP_PATH = Path('data/fixed_translation.json')
FIXED_MAP = {}

# ===================== 中文偵測與翻譯 =====================

def is_chinese(text):
    return bool(re.search('[\u4e00-\u9fff]', text))  # 判斷有沒有中文字

def translate_to_english(text):
    if not text or not text.strip() or not is_chinese(text):
        return text  # 空的或沒中文的直接跳過

    try:
        # deep_translator 有字數限制（通常 5000 字），需要分段處理
        max_length = 4500
        if len(text) > max_length:
            # 分段翻譯
            chunks = []
            sentences = text.split('\n')
            current_chunk = ""

            for sentence in sentences:
                if len(current_chunk) + len(sentence) + 1 <= max_length:
                    current_chunk += sentence + '\n'
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence + '\n'

            if current_chunk:
                chunks.append(current_chunk.strip())

            translated_chunks = []
            for chunk in chunks:
                if chunk:
                    translated = translator.translate(chunk)
                    translated_chunks.append(translated)
                    time.sleep(0.1)  # 避免請求過快

            return '\n'.join(translated_chunks)
        else:
            translated = translator.translate(text)
            return translated
    except Exception as e:
        print(f"翻譯錯誤: {text[:50]}... - {str(e)}")
        return text  # 翻譯失敗時返回原文

# ===================== 固定對照表讀取 =====================
def get_fixed_or_translator(text):
    stripped = text.strip()
    fixed = FIXED_MAP.get(stripped)
    if fixed:
        return fixed
    return translate_to_english(text)  # 沒有才呼叫翻譯器

# ===================== 判斷是否包含圖片 =====================
def has_picture(run):
    """檢查 run 是否包含圖片或繪圖對象"""
    # 檢查是否有 drawing 元素（圖片）
    if run._element.findall('.//' + qn('w:drawing')):
        return True
    # 檢查是否有 pict 元素（舊版圖片格式）
    if run._element.findall('.//' + qn('w:pict')):
        return True
    return False

# ===================== 安全清空段落文字（保留圖片）=====================
def clear_paragraph_text_keep_images(paragraph):
    """清空段落中的文字，但保留圖片"""
    for run in paragraph.runs:
        if not has_picture(run):
            run.text = ""
        # 包含圖片的 run 不清空

# ===================== 判斷是否為特殊格式 run =====================
def is_special_format(run):
    font = run.font
    if font.bold or font.italic or font.underline:
        return True
    if font.color and font.color.rgb and str(font.color.rgb) not in ('000000', None, ''):
        return True
    if font.highlight_color:
        return True
    # 如果包含圖片，也視為特殊格式
    if has_picture(run):
        return True
    return False

# ==================== 偵測開頭空格 ====================
def has_long_spaces_in_runs(paragraph):
    """檢測段落開頭是否有空格（修改為只要有1個空格就算）"""
    raw_parts = []
    for run in paragraph.runs:
        raw_parts.append(run.text)

    raw_text = ''.join(raw_parts)   # 把所有 run 拼回原本的樣子

    # 檢查開頭是否有空格或Tab
    if not raw_text:
        return False, 0

    # 計算開頭的空格數
    leading_spaces = len(raw_text) - len(raw_text.lstrip(' \t'))

    if leading_spaces > 0:
        # 將Tab轉換為4個空格計算
        space_count = 0
        for char in raw_text[:leading_spaces]:
            if char == '\t':
                space_count += 4
            elif char == ' ':
                space_count += 1

        return True, space_count

    return False, 0

# ==================== 偵測步驟編號 ====================
def get_step_number(paragraph):
    if isinstance(paragraph, str):
        raw = paragraph
    else:
        raw = ''.join(run.text for run in paragraph.runs)

    # 編號規則清單
    number_patterns = [
        r'^(\d+\.\d+\.\d+\.\d+)',       # 1.2.3.4
        r'^(\d+\.\d+.\d+)',              # 1.2.3
        r'^(\d+\.\d+)',                  # 1.2
        r'^(\d+\.)',                     # 1.
    ]

    for pattern in number_patterns:
        match = re.match(pattern, raw.lstrip())
        if match:
            num = match.group(1)
            return num.rstrip('.')  # 把最後的點去掉

    return ""  # 沒抓到就回傳空字串

# ==================== 檢測冒號格式 ====================
def check_colon_format(text):
    """
    檢查是否為冒號格式
    返回: (has_colon, has_content_after_colon, colon_part, content_part)
    """
    # 檢測中英文冒號
    colon_match = re.match(r'^([^:：]+)[：:](.*)$', text.strip())
    if colon_match:
        before_colon = colon_match.group(1).strip()
        after_colon = colon_match.group(2).strip()
        return True, bool(after_colon), before_colon, after_colon
    return False, False, "", ""

# ==================== 記錄長空格段落 =====================
def record_long_space_paragraph(paragraph, para_index=None):
    """
    修正原則：段落開頭有項目編號且後續段落的開頭有多空格的情況下合併
    如果後續段落遇到項目編號開頭，停止合併
    """
    global current_group, continuous_abnormal_groups
    has_abnormal, count = has_long_spaces_in_runs(paragraph)
    step_number = get_step_number(paragraph)

    if has_abnormal and para_index is not None:
        if step_number:
            # 有項目編號 → 結束舊組，開新組
            if current_group is not None:
                current_group["merged_text"] = merge_group_text(current_group["paragraphs"])
                continuous_abnormal_groups.append(current_group)
                current_group = None

            # 開新組
            current_group = {"group_id": len(continuous_abnormal_groups) + 1, "paragraphs": []}
            current_group["paragraphs"].append({
                "para_index": para_index,
                'para': paragraph,
                "full_text": paragraph.text.strip(),
                "space_count": count
            })
        else:
            # 沒有項目編號，但有多空格 → 添加到當前組（如果存在）
            if current_group is not None:
                current_group["paragraphs"].append({
                    "para_index": para_index,
                    'para': paragraph,
                    "full_text": paragraph.text.strip(),
                    "space_count": count
                })
            # 如果沒有當前組，說明這是單獨的縮排段落，不處理

# ==================== 合併同組縮排段落 ====================
def merge_group_text(paragraphs_list):
    if not paragraphs_list:
        return ""

    lines = []
    for p in paragraphs_list:
        text = p["full_text"]
        lines.append(text.lstrip())

    ## 判斷這整個 group 是否為「步驟編號型」
    is_step_group = any(get_step_number(p["full_text"]) for p in paragraphs_list)

    # === 合併邏輯 ===
    if is_step_group:
        result = lines[0]                    # 第一段保留
        for line in lines[1:]:
            result += line.lstrip()          # 後面全部直接接文字
        return result
    else:
        return " ".join(lines)               # 用空格連接

# ============================================================

def add_english_below(paragraph, english_text, font_size=None, font_name='Times New Roman', alignment=None):
    """
    在指定中文段落下方新增一行英文段落（中英對照用）
    """
    # 取得當前段落的位置
    parent = paragraph._element.getparent()

    # 檢查段落是否已被刪除
    if parent is None:
        return None

    para_index = parent.index(paragraph._element)

    # 創建新段落元素
    new_p = OxmlElement('w:p')

    # 插入到原段落後面
    parent.insert(para_index + 1, new_p)

    # 創建 paragraph 對象
    from docx.text.paragraph import Paragraph
    eng_para = Paragraph(new_p, paragraph._parent)

    # 設定對齊方式（繼承原段落或使用指定值）
    if alignment == 'center':
        eng_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        eng_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == 'justify':
        eng_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == 'left':
        eng_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif paragraph.alignment is not None:
        eng_para.alignment = paragraph.alignment
    else:
        eng_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 添加文字
    run = eng_para.add_run(english_text)

    # 字體設定
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 字體大小處理
    if font_size is not None:
        run.font.size = Pt(font_size)
    else:
        # 繼承上方中文段落大小
        if paragraph.runs and paragraph.runs[0].font.size:
            run.font.size = paragraph.runs[0].font.size
        else:
            run.font.size = Pt(11)  # 預設 11pt

    return eng_para

# ===================== 調整段落字體大小函式 =====================

def set_paragraph_font_size(paragraph, font_size_pt, target='all'):
    """
    改進版：以整個 paragraph 是否含中文來決定是否調整所有 run
    """
    if not paragraph.runs:
        return

    half_points = font_size_pt * 2

    # 先強制移除樣式綁定（頁首頁尾必備）
    pPr = paragraph._p.find(qn('w:pPr'))
    if pPr is not None:
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            pPr.remove(pStyle)

    # 整段 paragraph 文字合併判斷是否有中文
    full_text = paragraph.text
    paragraph_has_chinese = bool(re.search('[\u4e00-\u9fff]', full_text))

    # 根據 target 和整段情況決定是否調整所有 run
    should_adjust = (
        target == 'all' or
        (target == 'chinese' and paragraph_has_chinese) or
        (target == 'english' and not paragraph_has_chinese)
    )

    if not should_adjust:
        return

    # 統一調整這段所有 run
    for run in paragraph.runs:
        if not run.text.strip():
            continue

        run.font.size = Pt(font_size_pt)

        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._r.insert(0, rPr)

        for tag in ('w:sz', 'w:szCs'):
            elem = rPr.find(qn(tag))
            if elem is None:
                elem = OxmlElement(tag)
                rPr.insert(0, elem)
            elem.set(qn('w:val'), str(half_points))

# ===================== 翻譯段落函式（雙語版）=====================

def translate_paragraph_bilingual(paragraph, para_index=None):
    """
    雙語翻譯：保留中文，下方添加英文
    根據 format.txt 的格式要求處理
    """
    if not paragraph.text.strip() or not is_chinese(paragraph.text):
        return

    global continuous_abnormal_groups, translated_group_ids

    # ========= 判斷此 paragraph 是否屬於某個縮排群組 =========
    belonging_group = None
    is_first_para_in_group = False

    for group in continuous_abnormal_groups:
        if group["group_id"] in translated_group_ids:
            continue
        for item in group["paragraphs"]:
            if item["para_index"] == para_index:
                belonging_group = group
                is_first_para_in_group = (item["para_index"] == group["paragraphs"][0]["para_index"])
                break
        if belonging_group:
            break

    # ========= 如果是縮排群組的「非第一段」→ 跳過（已在第一段處理時刪除）=========
    if belonging_group and not is_first_para_in_group:
        # 跳過此段落（已在第一段處理時被刪除或將被刪除）
        return

    # ========= 如果是縮排群組的「第一段」→ 整組合併翻譯 =========
    if belonging_group and is_first_para_in_group:
        # 合併所有段落內容
        merged_chinese = merge_group_text(belonging_group["paragraphs"])

        # 提取編號（如果有）
        first_full_text = belonging_group["paragraphs"][0]["full_text"]
        step_number = get_step_number(first_full_text)

        # 保留原始縮排
        raw_text = paragraph.text
        leading_spaces = len(raw_text) - len(raw_text.lstrip())
        indent = " " * leading_spaces

        # 更新第一段落的中文內容為合併後的內容
        # 檢查是否有圖片
        has_images = any(has_picture(run) for run in paragraph.runs)

        if has_images:
            # 有圖片：只清空文字 run，保留圖片 run
            text_runs = [run for run in paragraph.runs if not has_picture(run)]
            for run in text_runs:
                run.text = ""
            # 在最後添加合併後的文字
            paragraph.add_run(merged_chinese)
        else:
            # 無圖片：清空所有 run 並重寫
            for run in paragraph.runs:
                run.text = ""
            # 寫入合併後的中文
            final_chinese = indent + merged_chinese
            paragraph.add_run(final_chinese)

        # 去掉編號翻譯
        if step_number:
            content_start = merged_chinese.find(step_number) + len(step_number)
            while content_start < len(merged_chinese) and merged_chinese[content_start] in '. :：\u3000\t ':
                content_start += 1
            pure_content = merged_chinese[content_start:]
        else:
            pure_content = merged_chinese

        # 翻譯
        translated_full = translate_to_english(pure_content)

        # 組合英文文字（不包含編號，保留縮進）
        final_english = indent + translated_full.strip()

        # 在第一段後添加英文段落
        add_english_below(paragraph, final_english)

        # 標記其他段落為空（稍後統一刪除）
        for item in belonging_group["paragraphs"][1:]:
            # 清空後續段落的內容（保留圖片）
            clear_paragraph_text_keep_images(item["para"])

        translated_group_ids.add(belonging_group["group_id"])
        return

    # ========= 非縮排群組的普通段落 =========
    raw_text = paragraph.text
    stripped_text = raw_text.strip()

    # 檢查冒號格式
    has_colon, has_content, colon_part, _ = check_colon_format(stripped_text)

    if has_colon and not has_content:
        # 格式一：冒號後無字 → 同行括號添加英文
        # 例：權責： → 權責(Authority and Responsibility):
        translated = translate_to_english(colon_part)

        # 檢查是否有圖片
        has_images = any(has_picture(run) for run in paragraph.runs)

        # 判斷原始冒號類型
        original_colon = '：' if '：' in raw_text else ':'
        leading_spaces = len(raw_text) - len(raw_text.lstrip())
        indent = " " * leading_spaces
        new_text = f"{indent}{colon_part}({translated}){original_colon}"

        if has_images:
            # 有圖片：只清空文字 run，保留圖片
            for run in paragraph.runs:
                if not has_picture(run):
                    run.text = ""
            paragraph.add_run(new_text)
        else:
            # 無圖片：清空所有 run 並重寫
            for run in paragraph.runs:
                run.text = ""
            paragraph.add_run(new_text)

    elif has_colon and has_content:
        # 格式二：冒號後有字 → 換行添加英文
        # 保持中文不變，下方添加英文

        # 檢查是否有步驟編號
        step_number = get_step_number(stripped_text)

        if step_number:
            # 有編號，去掉編號部分翻譯
            content_start = stripped_text.find(step_number) + len(step_number)
            while content_start < len(stripped_text) and stripped_text[content_start] in '. :：\u3000\t ':
                content_start += 1
            pure_content = stripped_text[content_start:]
            translated = translate_to_english(pure_content)

            # 保留縮排（不包含編號）
            leading_spaces = len(raw_text) - len(raw_text.lstrip())
            indent = " " * leading_spaces
            final_english = f"{indent}{translated}"
        else:
            # 翻譯整段
            translated = translate_to_english(stripped_text)

            # 保留縮排
            leading_spaces = len(raw_text) - len(raw_text.lstrip())
            indent = " " * leading_spaces
            final_english = f"{indent}{translated}"

        # 在下方添加英文段落
        add_english_below(paragraph, final_english)

    else:
        # 格式三：沒有冒號 → 換行添加英文
        step_number = get_step_number(stripped_text)

        if step_number:
            # 有編號
            content_start = stripped_text.find(step_number) + len(step_number)
            while content_start < len(stripped_text) and stripped_text[content_start] in '. :：\u3000\t ':
                content_start += 1
            pure_content = stripped_text[content_start:]
            translated = translate_to_english(pure_content)

            leading_spaces = len(raw_text) - len(raw_text.lstrip())
            indent = " " * leading_spaces
            final_english = f"{indent}{translated}"
        else:
            # 無編號
            translated = translate_to_english(stripped_text)

            leading_spaces = len(raw_text) - len(raw_text.lstrip())
            indent = " " * leading_spaces
            final_english = f"{indent}{translated}"

        # 在下方添加英文段落
        add_english_below(paragraph, final_english)

# ===================== 翻譯表格函式（雙語版）=====================

def translate_table_bilingual(table):
    """
    表格雙語翻譯
    根據 format.txt 要求：如果有編號，所有中文段落保留，譯文在最後合併添加
    """
    for row in table.rows:
        for cell in row.cells:
            # 收集所有段落
            paragraphs = [p for p in cell.paragraphs if p.text.strip() and is_chinese(p.text)]

            if not paragraphs:
                continue

            # 檢查是否有編號格式
            has_numbers = any(get_step_number(p.text) for p in paragraphs)

            if has_numbers and len(paragraphs) > 1:
                # 表格內有多個編號段落 → 合併翻譯，添加在最後
                all_chinese = []
                all_numbers = []

                for p in paragraphs:
                    stripped = p.text.strip()
                    step_num = get_step_number(stripped)

                    if step_num:
                        content_start = stripped.find(step_num) + len(step_num)
                        while content_start < len(stripped) and stripped[content_start] in '. :：\u3000\t ':
                            content_start += 1
                        pure_content = stripped[content_start:]
                        all_chinese.append(pure_content)
                        all_numbers.append(step_num)
                    else:
                        all_chinese.append(stripped)
                        all_numbers.append("")

                # 合併翻譯
                merged_chinese = " ".join(all_chinese)
                translated_full = translate_to_english(merged_chinese)

                # 將譯文按編號拆分（簡化處理：按句號拆分）
                translated_parts = translated_full.split('.')
                translated_parts = [p.strip() for p in translated_parts if p.strip()]

                # 組合英文編號段落（表格內保留編號）
                english_lines = []
                for i, num in enumerate(all_numbers):
                    if i < len(translated_parts):
                        if num:
                            english_lines.append(f"{num}.{translated_parts[i]}")
                        else:
                            english_lines.append(translated_parts[i])
                    else:
                        break

                # 如果有剩餘的譯文，加到最後
                if len(translated_parts) > len(all_numbers):
                    english_lines.extend(translated_parts[len(all_numbers):])

                # 在最後一個段落後添加英文
                last_para = paragraphs[-1]
                combined_english = "\n".join(english_lines)
                add_english_below(last_para, combined_english)

            else:
                # 單個段落或無編號 → 正常處理
                for p in paragraphs:
                    translate_paragraph_bilingual(p)

# ===================== 翻譯頁首頁尾函式 =====================

def translate_header_footer_full(doc):
    """完全通用版：翻譯頁首頁尾的「文字 + 表格 + 表格內文字」"""
    for section in doc.sections:

        # 所有可能的頁首
        headers = [
            section.header,
            section.first_page_header,
            section.even_page_header
        ]
        # 所有可能的頁尾
        footers = [
            section.footer,
            section.first_page_footer,
            section.even_page_footer
        ]

        for hf in headers + footers:
            if not hf or hf.is_linked_to_previous:
                continue

            # 1. 翻頁首頁尾內的一般段落
            for para in hf.paragraphs:
                if para.text.strip() and is_chinese(para.text):
                    set_paragraph_font_size(para, 10, 'chinese')
                    print(f"讀取: {para.text}")
                    translated = get_fixed_or_translator(para.text)
                    if translated:
                        add_english_below(para, translated, font_size=6, alignment='center')

            # 2. 翻頁首頁尾內的「表格」
            for table in hf.tables:
                processed_texts = set()
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            full_text = para.text.strip()
                            if para.text.strip() and is_chinese(para.text):
                                set_paragraph_font_size(para, 10, 'chinese')
                                if full_text in processed_texts:
                                    continue

                                print(f"讀取: {para.text}")
                                translated = get_fixed_or_translator(para.text)
                                if translated:
                                    add_english_below(para, translated, font_size=8, alignment='center')
                                    processed_texts.add(full_text)

# ===================== 翻譯流程圖文字函式 =====================

def translate_textboxes_in_doc(doc):
    body = doc.element.body
    if body is None:
        return

    textboxes = body.findall('.//' + qn('w:txbxContent'))

    if not textboxes:
        return

    for textbox in textboxes:
        # 翻譯所有文字
        text_elements = textbox.findall('.//' + qn('w:t'))
        for text_elem in text_elements:
            if text_elem.text and text_elem.text.strip():
                original_text = text_elem.text
                translated_text = translate_to_english(original_text)
                if translated_text:
                    text_elem.text = translated_text

                if is_chinese(original_text):
                    half_points = 11
                    # 調整字體大小
                    for r in textbox.findall('.//' + qn('w:r')):
                        rPr = r.find(qn('w:rPr'))
                        if rPr is None:
                            rPr = OxmlElement('w:rPr')
                            r.append(rPr)
                        sz = rPr.find(qn('w:sz'))
                        if sz is None:
                            sz = OxmlElement('w:sz')
                            rPr.append(sz)
                        sz.set(qn('w:val'), str(half_points))

                    # 調整行高
                    for p in textbox.findall('.//' + qn('w:p')):
                        pPr = p.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            p.insert(0, pPr)
                        spacing = pPr.find(qn('w:spacing'))
                        if spacing is None:
                            spacing = OxmlElement('w:spacing')
                            pPr.append(spacing)
                        spacing.set(qn('w:line'), '130')
                        spacing.set(qn('w:lineRule'), 'exact')

                    # 調整對齊
                    for p in textbox.findall('.//' + qn('w:p')):
                        pPr = p.find(qn('w:pPr'))
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            p.insert(0, pPr)

                        jc = pPr.find(qn('w:jc'))
                        if jc is None:
                            jc = OxmlElement('w:jc')
                            pPr.append(jc)
                        jc.set(qn('w:val'), 'center')

# ===================== 縮小表格內英文文字函式 =====================
def shrink_table_english_font(table, ratio=0.82):
    """
    縮小表格內「純英文」文字大小
    """
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text = paragraph.text.strip()
                if not text or any('\u4e00' <= c <= '\u9fff' for c in text):
                    continue

                for run in paragraph.runs:
                    if not run.text.strip():
                        continue

                    r = run._r
                    rPr = r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = OxmlElement('w:rPr')
                        r.insert(0, rPr)

                    # 處理 w:sz
                    sz = rPr.find(qn('w:sz'))
                    if sz is None:
                        sz = OxmlElement('w:sz')
                        rPr.insert(0, sz)

                    if sz.get(qn('w:val')):
                        current = int(sz.get(qn('w:val')))
                        new_val = max(20, int(current * ratio))
                    else:
                        new_val = max(20, int(22 * ratio))

                    sz.set(qn('w:val'), str(new_val))

                    # 處理 w:szCs
                    szCs = rPr.find(qn('w:szCs'))
                    if szCs is None:
                        szCs = OxmlElement('w:szCs')
                        rPr.insert(0, szCs)
                    szCs.set(qn('w:val'), str(new_val))

                    # 處理 w:szFarEast
                    szFarEast = rPr.find(qn('w:szFarEast'))
                    if szFarEast is None:
                        szFarEast = OxmlElement('w:szFarEast')
                        rPr.insert(0, szFarEast)
                    szFarEast.set(qn('w:val'), str(new_val))

# ===================== 強制 Times New Roman 字體函式 =====================
def force_times_new_roman(doc):
    """強制全文件所有文字改成 Times New Roman（只針對英文，中文保留原始字體）"""
    def _set_font(run):
        r = run._r
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)

        # 判斷 run 內的文字是否包含中文
        text = run.text.strip()
        has_chinese = bool(re.search('[\u4e00-\u9fff]', text)) if text else False

        if not has_chinese and text:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rFonts.set(qn('w:cs'), 'Times New Roman')

    # 正文
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_font(run)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_font(run)

    # 頁首頁尾
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            if hf:
                for paragraph in hf.paragraphs:
                    for run in paragraph.runs:
                        _set_font(run)
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    _set_font(run)

    # 流程圖文字框
    for textbox in doc.element.body.findall('.//' + qn('w:txbxContent')):
        for run in textbox.findall('.//' + qn('w:r')):
            rPr = run.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run.insert(0, rPr)

            # 同樣判斷中文
            text = run.text.strip()
            has_chinese = bool(re.search('[\u4e00-\u9fff]', text)) if text else False

            if not has_chinese and text:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:ascii'), 'Times New Roman')
                rFonts.set(qn('w:hAnsi'), 'Times New Roman')
                rFonts.set(qn('w:cs'), 'Times New Roman')


# ===================== 清理空段落函式 =====================
def remove_empty_paragraphs(doc):
    """刪除文檔中的空白段落（保留包含圖片的段落）"""
    paragraphs_to_remove = []

    for paragraph in doc.paragraphs:
        # 如果段落文字為空，但檢查是否包含圖片
        if not paragraph.text.strip():
            # 檢查是否有圖片
            has_images = any(has_picture(run) for run in paragraph.runs)
            if not has_images:
                # 沒有圖片才刪除
                paragraphs_to_remove.append(paragraph)

    for para in paragraphs_to_remove:
        p_element = para._element
        parent = p_element.getparent()
        if parent is not None:
            parent.remove(p_element)

# ==================== 主翻譯函式（雙語版）====================
def translate_document(input_file, output_file):
    global continuous_abnormal_groups, current_group, translated_group_ids, FIXED_MAP, FIXED_MAP_PATH
    continuous_abnormal_groups = []
    current_group = None
    translated_group_ids = set()

    # 載入固定翻譯對照表
    if FIXED_MAP_PATH.exists():
        with FIXED_MAP_PATH.open('r', encoding='utf-8') as f:
            FIXED_MAP = json.load(f)
    print(f"載入固定翻譯對照表，共 {len(FIXED_MAP)} 筆資料")

    print(f"="*90)
    print(f"載入檔案：{input_file}\n")
    doc = Document(input_file)

    print("開始正文段落合併...")
    # 先記錄全部的長空格段落跟合併
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            if not paragraph.text.strip() or not is_chinese(paragraph.text):
                continue
            record_long_space_paragraph(paragraph, para_index=i)

    # 結束最後一組
    if current_group is not None:
        current_group["merged_text"] = merge_group_text(current_group["paragraphs"])
        continuous_abnormal_groups.append(current_group)

    print("開始翻譯正文段落（雙語模式）...")
    # 進行雙語翻譯
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            translate_paragraph_bilingual(paragraph, para_index=i)

    print("\n開始翻譯表格內容（雙語模式）...")
    for table in doc.tables:
        translate_table_bilingual(table)

    print("\n開始翻譯頁首頁尾...")
    translate_header_footer_full(doc)

    print(f"翻譯流程圖.....\n")
    translate_textboxes_in_doc(doc)

    print("開始縮小表格內英文字體（82%）...")
    for table in doc.tables:
        shrink_table_english_font(table, ratio=0.82)

    # 頁首頁尾表格也要
    for section in doc.sections:
        for hf in (section.header, section.footer,
                section.first_page_header, section.first_page_footer,
                section.even_page_header, section.even_page_footer):
            if hf and hf.tables:
                for t in hf.tables:
                    shrink_table_english_font(t, ratio=0.82)

    print("強制全文件字體為 Times New Roman...")
    force_times_new_roman(doc)

    print("清理空白段落...")
    remove_empty_paragraphs(doc)

    print(f"="*90)
    print(f"儲存翻譯結果 → {output_file}")
    doc.save(output_file)
    print("翻譯完成！")

# ==================== 一鍵執行 ====================
if __name__ == "__main__":

    print(f'多語文檔轉譯專案 (使用 Deep Translator - 雙語版本)\n')

    # ===================== 主程式執行 =====================
    test_file = Path(("document_cn.docx"))


    output_file = Path(("out_bilingual.docx"))
    print(f"輸出檔案：{output_file}\n")

    start_time = time.time()

    translate_document(test_file, output_file)

    total_time = time.time() - start_time
    print(f"總耗時：{total_time:.2f} 秒（{total_time/60:.2f} 分鐘）")

    from datetime import datetime
    with open("翻譯時間紀錄.log", "a", encoding="utf-8") as f:
        f.write(f"{datetime.now():%Y-%m-%d %H:%M:%S} | [DeepTranslator-Bilingual] | 總耗時: {total_time/60:.2f} 分鐘\n")

