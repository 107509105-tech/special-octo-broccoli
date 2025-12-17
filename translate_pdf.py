#!/usr/bin/env python3
"""
PDF 轉換與翻譯腳本
將 PDF 轉換為 DOCX，然後翻譯成繁體中文
"""

from pdf2docx import Converter
from docx import Document
from googletrans import Translator
import time
import sys

def pdf_to_docx(pdf_path, docx_path):
    """將 PDF 轉換為 DOCX"""
    print(f"正在將 {pdf_path} 轉換為 DOCX...")
    cv = Converter(pdf_path)
    cv.convert(docx_path, start=0, end=None)
    cv.close()
    print(f"轉換完成，保存為 {docx_path}")

def translate_docx(input_docx, output_docx):
    """翻譯 DOCX 文件內容為繁體中文"""
    print(f"正在讀取 {input_docx}...")
    doc = Document(input_docx)
    translator = Translator()

    print("開始翻譯...")
    total_paragraphs = len(doc.paragraphs)

    # 翻譯段落
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            try:
                # 保留原始格式
                original_text = paragraph.text
                print(f"翻譯段落 {i+1}/{total_paragraphs}: {original_text[:50]}...")

                # 翻譯為繁體中文
                translated = translator.translate(original_text, src='auto', dest='zh-tw')
                paragraph.text = translated.text

                # 避免 API 請求過快
                time.sleep(0.5)

            except Exception as e:
                print(f"翻譯段落 {i+1} 時發生錯誤: {e}")
                continue

    # 翻譯表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    try:
                        original_text = cell.text
                        print(f"翻譯表格內容: {original_text[:30]}...")
                        translated = translator.translate(original_text, src='auto', dest='zh-tw')
                        cell.text = translated.text
                        time.sleep(0.5)
                    except Exception as e:
                        print(f"翻譯表格時發生錯誤: {e}")
                        continue

    # 保存翻譯後的文件
    print(f"保存翻譯後的文件至 {output_docx}...")
    doc.save(output_docx)
    print("翻譯完成！")

def main():
    # 設定文件路徑
    pdf_path = "ysm20r.pdf"
    temp_docx = "ysm20r_temp.docx"
    output_docx = "ysm20r_translated.docx"

    try:
        # 步驟 1: PDF 轉 DOCX
        pdf_to_docx(pdf_path, temp_docx)

        # 步驟 2: 翻譯 DOCX
        translate_docx(temp_docx, output_docx)

        print(f"\n✓ 完成！翻譯後的文件: {output_docx}")

    except Exception as e:
        print(f"\n✗ 發生錯誤: {e}")
        sys.exit(1)

if __name__ == "__main__":
    main()
